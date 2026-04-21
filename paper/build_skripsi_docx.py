"""Generate FULL/expanded skripsi-style Word document (Bab 1-5) in Indonesian."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "paper"
OUT_DIR.mkdir(exist_ok=True)

evaluate = json.loads((ROOT / "output/results/evaluate.json").read_text())
comparison = json.loads((ROOT / "output/results/comparison_summary.json").read_text())

doc = Document()

for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def H1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"


def H3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"


def P(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"


def BULLET(items, numbered=False):
    style_name = "List Number" if numbered else "List Bullet"
    for it in items:
        try:
            p = doc.add_paragraph(it, style=style_name)
        except KeyError:
            prefix = "" if numbered else "• "
            p = doc.add_paragraph(prefix + it)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"


def CAPTION(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"


def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)


def TABLE(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
    doc.add_paragraph()


# ============ COVER ============
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("ANALISIS KOMPARATIF METODE PERINGKASAN TEKS\nEKSTRAKTIF DAN ABSTRAKTIF UNTUK\nDOKUMEN AKADEMIK BERBAHASA INDONESIA")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

for _ in range(3):
    doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Laporan Penelitian / Skripsi\n\nDisusun oleh:\n[Nama Mahasiswa]\n[NIM]\n\nProgram Studi Informatika\nFakultas Teknik\n[Nama Universitas]\n2026")
r.font.size = Pt(12)
r.font.name = "Times New Roman"

doc.add_page_break()

# ============================================================
# BAB 1 — PENDAHULUAN
# ============================================================
H1("BAB 1\nPendahuluan")

H2("1.1 Latar Belakang")
P("Era digital ditandai dengan ledakan informasi yang sangat cepat. Setiap harinya, ratusan ribu artikel ilmiah, jurnal, prosiding, dan dokumen akademik dipublikasikan di seluruh dunia. Di Indonesia, jumlah publikasi ilmiah berbahasa Indonesia juga terus meningkat seiring dengan tumbuhnya jurnal nasional terindeks SINTA, repositori institusi perguruan tinggi, serta meningkatnya budaya menulis ilmiah di kalangan akademisi. Pertumbuhan ini, meskipun positif, justru menimbulkan tantangan baru: bagaimana seorang peneliti, dosen, atau mahasiswa dapat memilah, membaca, dan memahami volume informasi yang sangat besar dalam waktu yang terbatas.")
P("Berdasarkan data dari dataset penelitian ini, sebuah dokumen akademik tunggal memiliki rata-rata panjang sebesar 25.267 karakter atau setara dengan kurang lebih 4.000 hingga 5.000 kata, dengan dokumen terpanjang mencapai 57.745 karakter. Apabila seseorang ingin meninjau seratus dokumen sekaligus untuk kebutuhan studi literatur (literature review), maka ia harus membaca lebih dari dua setengah juta karakter atau setara dengan ratusan halaman teks. Aktivitas ini tentu sangat menyita waktu dan tidak efisien, terlebih jika tidak semua dokumen relevan dengan topik penelitian yang sedang dikerjakan.")
P("Untuk mengatasi masalah ini, peringkasan teks otomatis (automatic text summarization) menjadi salah satu solusi yang sangat menarik dalam ranah Natural Language Processing (NLP). Peringkasan teks otomatis adalah proses menghasilkan ringkasan yang lebih singkat dari sebuah dokumen sumber dengan tetap mempertahankan informasi yang paling penting. Hasil ringkasan dapat membantu pembaca memahami inti dokumen secara cepat tanpa harus membaca keseluruhan teks.")
P("Pada literatur NLP modern, peringkasan teks dibagi menjadi dua paradigma utama, yaitu pendekatan ekstraktif (extractive summarization) dan pendekatan abstraktif (abstractive summarization). Pendekatan ekstraktif bekerja dengan cara memilih kalimat-kalimat penting dari dokumen sumber dan menggabungkannya menjadi ringkasan baru. Karena kalimat-kalimat tersebut diambil langsung dari dokumen asli, ringkasan ekstraktif cenderung memiliki struktur yang gramatikal dan konsisten dengan teks sumber. Beberapa algoritma populer untuk pendekatan ini adalah TF-IDF, Latent Semantic Analysis (LSA), dan TextRank yang berbasis graph PageRank.")
P("Sebaliknya, pendekatan abstraktif berusaha menghasilkan kalimat baru yang tidak harus muncul secara persis di dokumen sumber. Pendekatan ini meniru cara manusia meringkas, yaitu dengan memahami makna teks terlebih dahulu kemudian menulis ulang dalam bahasa yang lebih singkat. Pendekatan ini menjadi sangat populer setelah munculnya arsitektur Transformer dan pre-trained language model seperti BART, T5, Pegasus, dan varian multilingualnya seperti mBART dan mT5.")
P("Sebagian besar penelitian peringkasan teks selama ini berfokus pada bahasa Inggris karena ketersediaan dataset yang melimpah dan dukungan model pre-trained yang sangat banyak. Sayangnya, untuk bahasa Indonesia, terutama pada domain dokumen akademik, riset komparatif yang sistematis dengan metrik evaluasi standar masih sangat terbatas. Bahasa Indonesia memiliki karakteristik morfologi aglutinatif yang unik, di mana satu kata dasar dapat membentuk banyak kata turunan melalui afiksasi (awalan, sisipan, akhiran, dan konfiks) seperti me-, di-, ber-, ter-, ke-, -an, -kan, -i, dan kombinasinya. Hal ini menyebabkan ukuran kosakata bahasa Indonesia menjadi sangat besar dan menjadi tantangan tersendiri bagi sistem NLP, terutama dalam tahap stemming dan tokenization.")
P("Di sisi lain, hadirnya Large Language Model (LLM) seperti ChatGPT, Claude, Gemini, dan model open-source lainnya turut mengubah lanskap penelitian peringkasan teks. LLM mampu menghasilkan ringkasan dengan kualitas yang sangat tinggi, namun penggunaannya bergantung pada koneksi internet, biaya API, dan tidak selalu transparan dalam hal proses internal yang digunakannya. Karena itu, perlu dilakukan studi komparatif yang membandingkan metode klasik (ekstraktif), metode neural berbasis pre-trained model (abstraktif), dan output LLM untuk memberikan gambaran utuh kepada para pengguna.")
P("Penelitian ini dilakukan untuk menjawab kekosongan tersebut. Penulis melakukan implementasi dan evaluasi dua metode peringkasan teks pada 100 dokumen akademik berbahasa Indonesia, yaitu metode ekstraktif berbasis TF-IDF + TextRank dan metode abstraktif berbasis model mT5 multilingual hasil fine-tuning XL-Sum. Selain itu, penulis juga membandingkan hasil keduanya terhadap ringkasan yang dihasilkan oleh LLM pada subset 10 dokumen sebagai third baseline. Sebagai bentuk implementasi praktis, penulis juga membangun aplikasi web interaktif berbasis Flask yang memungkinkan pengguna menjalankan keseluruhan pipeline mulai dari preprocessing hingga evaluasi melalui antarmuka yang user-friendly.")

H2("1.2 Identifikasi Masalah")
P("Berdasarkan latar belakang di atas, dapat diidentifikasi beberapa permasalahan sebagai berikut:")
BULLET([
    "Volume dokumen akademik berbahasa Indonesia tumbuh sangat cepat sehingga sulit dibaca secara manual.",
    "Riset komparatif metode ekstraktif vs abstraktif untuk dokumen akademik berbahasa Indonesia masih jarang.",
    "Model pre-trained khusus bahasa Indonesia untuk tugas summarization masih terbatas, sehingga umumnya peneliti menggunakan model multilingual seperti mT5.",
    "Belum banyak tersedia tool open-source berbahasa Indonesia yang menyediakan pipeline lengkap mulai preprocessing hingga evaluasi.",
    "Belum ada studi yang sistematis dalam menghubungkan karakteristik dokumen (panjang, jumlah kalimat, rasio kompresi) dengan kualitas ringkasan yang dihasilkan.",
])

H2("1.3 Rumusan Masalah")
P("Berdasarkan identifikasi masalah, dirumuskan permasalahan penelitian sebagai berikut:")
BULLET([
    "Bagaimana performa metode ekstraktif berbasis TF-IDF + TextRank dalam meringkas dokumen akademik berbahasa Indonesia?",
    "Bagaimana performa metode abstraktif berbasis mT5 multilingual (XL-Sum) pada dataset yang sama?",
    "Metode mana yang lebih unggul berdasarkan metrik ROUGE-1, ROUGE-2, dan ROUGE-L (Precision, Recall, dan F1)?",
    "Faktor-faktor apa, seperti panjang dokumen, jumlah kalimat, dan rasio kompresi, yang memengaruhi kualitas ringkasan per dokumen?",
    "Bagaimana hasil ringkasan kedua metode ketika dibandingkan dengan ringkasan yang dihasilkan oleh Large Language Model (LLM) sebagai referensi tambahan?",
    "Bagaimana pipeline peringkasan teks tersebut dapat diintegrasikan ke dalam aplikasi web interaktif yang dapat digunakan oleh pengguna umum?",
], numbered=True)

H2("1.4 Tujuan Penelitian")
P("Penelitian ini memiliki tujuan-tujuan sebagai berikut:")
BULLET([
    "Mengimplementasikan metode peringkasan ekstraktif berbasis TF-IDF dan TextRank untuk teks akademik Indonesia.",
    "Mengimplementasikan metode peringkasan abstraktif menggunakan model pre-trained csebuetnlp/mT5_multilingual_XLSum.",
    "Mengevaluasi kedua metode menggunakan metrik ROUGE pada 100 dokumen akademik berbahasa Indonesia.",
    "Menganalisis faktor-faktor per dokumen yang memengaruhi kualitas ringkasan.",
    "Membandingkan hasil ringkasan kedua metode dengan ringkasan yang dihasilkan oleh LLM pada subset 10 dokumen.",
    "Mengembangkan aplikasi web interaktif berbasis Flask sebagai bentuk implementasi praktis dari hasil penelitian.",
], numbered=True)

H2("1.5 Batasan Masalah")
P("Agar penelitian ini lebih terfokus dan terukur, ditetapkan batasan-batasan sebagai berikut:")
BULLET([
    "Bahasa yang digunakan adalah bahasa Indonesia.",
    "Domain dokumen yang diteliti adalah dokumen akademik (jurnal, skripsi, atau artikel ilmiah).",
    "Jumlah dokumen pada dataset adalah 100 dokumen.",
    "Model abstraktif yang digunakan tidak melalui fine-tuning karena keterbatasan sumber daya GPU; inferensi dilakukan menggunakan CPU AMD EPYC 7763.",
    "Metrik evaluasi yang digunakan terbatas pada ROUGE-1, ROUGE-2, dan ROUGE-L (Precision, Recall, F1).",
    "Aplikasi web yang dibangun bersifat lokal (Flask + Gunicorn + Nginx untuk deploy), tanpa autentikasi pengguna.",
    "Penelitian tidak melakukan evaluasi manusia (human evaluation) secara formal karena keterbatasan waktu dan responden.",
])

H2("1.6 Manfaat Penelitian")
H3("1.6.1 Manfaat Akademis")
P("Penelitian ini memberikan kontribusi akademis berupa baseline komparatif metode peringkasan teks pada domain akademik berbahasa Indonesia. Hasil penelitian ini dapat dijadikan acuan oleh peneliti lain yang ingin mengembangkan metode peringkasan teks Indonesia, baik dengan pendekatan klasik maupun pendekatan deep learning. Selain itu, dataset, kode sumber, dan dokumentasi penelitian dapat dijadikan bahan referensi pembelajaran NLP berbahasa Indonesia.")
H3("1.6.2 Manfaat Praktis")
P("Secara praktis, aplikasi web yang dihasilkan dapat langsung digunakan oleh akademisi, mahasiswa, peneliti, dan praktisi untuk meringkas dokumen-dokumen panjang berbahasa Indonesia. Aplikasi ini mendukung input dalam berbagai format (CSV, PDF, dan teks manual), sehingga dapat menjawab kebutuhan riil di lapangan.")
H3("1.6.3 Manfaat Metodologis")
P("Pipeline preprocessing bahasa Indonesia (case folding, cleaning, tokenization, stopword removal Sastrawi, stemming Sastrawi) yang dibangun dalam penelitian ini dapat dipakai ulang untuk tugas-tugas NLP Indonesia lainnya, seperti klasifikasi teks, sentiment analysis, dan information retrieval.")

H2("1.7 Metodologi Penelitian Singkat")
P("Penelitian ini dilakukan dengan pendekatan kuantitatif eksperimen. Tahapan penelitian secara garis besar meliputi: studi pustaka, pengumpulan dataset, preprocessing, implementasi metode ekstraktif, implementasi metode abstraktif, evaluasi ROUGE, perbandingan dengan LLM, pengembangan aplikasi web, dan terakhir analisis serta penulisan laporan. Detail metodologi disampaikan pada Bab 3.")

H2("1.8 Sistematika Penulisan")
P("Laporan penelitian ini disusun ke dalam lima bab sebagai berikut:")
BULLET([
    "Bab 1 Pendahuluan — berisi latar belakang, identifikasi masalah, rumusan masalah, tujuan, batasan, manfaat, metodologi singkat, dan sistematika penulisan.",
    "Bab 2 Tinjauan Pustaka — membahas konsep dasar Natural Language Processing, peringkasan teks ekstraktif dan abstraktif, TF-IDF, cosine similarity, PageRank/TextRank, arsitektur Transformer, model mT5 dan XL-Sum, metrik ROUGE, serta penelitian-penelitian terkait.",
    "Bab 3 Metodologi Penelitian — menjelaskan alur penelitian, dataset, pipeline preprocessing, implementasi metode ekstraktif dan abstraktif, prosedur evaluasi, perbandingan dengan LLM, pengembangan aplikasi web, dan spesifikasi sistem.",
    "Bab 4 Hasil dan Pembahasan — menyajikan hasil preprocessing, hasil ringkasan, hasil evaluasi ROUGE pada dataset penuh dan subset, analisis trade-off precision-recall, analisis faktor per dokumen, hasil aplikasi web, serta diskusi keterbatasan.",
    "Bab 5 Kesimpulan dan Saran — menyimpulkan temuan utama penelitian dan memberikan saran untuk penelitian lanjutan.",
])

doc.add_page_break()

# ============================================================
# BAB 2 — TINJAUAN PUSTAKA
# ============================================================
H1("BAB 2\nTinjauan Pustaka")

H2("2.1 Natural Language Processing (NLP)")
P("Natural Language Processing (NLP) adalah salah satu cabang ilmu Artificial Intelligence (AI) yang berfokus pada interaksi antara komputer dan bahasa manusia (natural language). Tujuan utama NLP adalah memungkinkan komputer untuk membaca, memahami, menganalisis, dan menghasilkan teks atau ucapan manusia secara bermakna. Beberapa contoh aplikasi NLP yang umum dijumpai adalah search engine, mesin penerjemah (machine translation), chatbot, voice assistant, sentiment analysis, named entity recognition, dan text summarization.")
P("Secara umum, sebuah pipeline NLP terdiri dari beberapa tahapan dasar yang saling berurutan dan saling melengkapi. Tahapan-tahapan ini bertujuan untuk memproses teks mentah menjadi representasi yang lebih terstruktur dan dapat dimanfaatkan oleh model machine learning di tahap selanjutnya.")

H3("2.1.1 Case Folding")
P("Case folding adalah proses penyeragaman huruf besar dan kecil dalam teks. Pada umumnya, semua huruf dikonversi menjadi huruf kecil (lowercase) untuk memastikan bahwa kata seperti 'Indonesia', 'INDONESIA', dan 'indonesia' dianggap identik oleh sistem.")

H3("2.1.2 Cleaning")
P("Cleaning meliputi penghapusan karakter-karakter yang tidak relevan, seperti tanda baca yang berlebih, simbol, angka (jika tidak dibutuhkan), URL, mention, hashtag, dan whitespace berlebih. Tujuannya adalah agar teks lebih bersih dan tidak menimbulkan noise pada tahap analisis berikutnya.")

H3("2.1.3 Sentence Tokenization")
P("Sentence tokenization adalah proses memecah paragraf atau dokumen menjadi unit-unit kalimat. Library yang sering dipakai adalah NLTK dengan fungsi sent_tokenize. Untuk bahasa Indonesia, NLTK juga dapat dipakai karena pemisah kalimat (titik, tanda tanya, tanda seru) bersifat universal.")

H3("2.1.4 Word Tokenization")
P("Setelah dipecah menjadi kalimat, setiap kalimat dipecah lagi menjadi kata-kata (token). Word tokenization dapat dilakukan dengan whitespace tokenizer, regex tokenizer, atau menggunakan library NLTK (word_tokenize).")

H3("2.1.5 Stopword Removal")
P("Stopword adalah kata-kata yang sering muncul dalam suatu bahasa tetapi tidak memberikan kontribusi semantik yang signifikan, seperti 'yang', 'di', 'ke', 'dari', 'dan', 'atau', 'adalah', 'untuk'. Penghapusan stopword bertujuan untuk mengurangi noise dan memperkecil dimensi vektor representasi. Daftar stopword bahasa Indonesia yang umum digunakan adalah daftar yang disediakan oleh library Sastrawi.")

H3("2.1.6 Stemming")
P("Stemming adalah proses pemotongan afiksasi (awalan, sisipan, akhiran, konfiks) untuk mendapatkan bentuk dasar (root) dari sebuah kata. Untuk bahasa Indonesia, library Sastrawi menjadi standar de-facto. Sastrawi mengimplementasikan algoritma Nazief–Adriani yang dirancang khusus untuk mengakomodasi morfologi bahasa Indonesia. Contoh: 'mengembangkan' menjadi 'kembang', 'menyimpulkan' menjadi 'simpul', 'pembelajaran' menjadi 'ajar'.")

H2("2.2 Peringkasan Teks (Text Summarization)")
P("Peringkasan teks adalah proses menghasilkan versi singkat dari sebuah teks atau dokumen sambil tetap mempertahankan informasi yang paling penting dan tidak menghilangkan makna utama dari teks asli. Peringkasan teks dapat dilakukan secara manual oleh manusia (human summary) atau secara otomatis oleh komputer (automatic summarization).")

H3("2.2.1 Klasifikasi Peringkasan Teks")
P("Peringkasan teks dapat diklasifikasikan berdasarkan beberapa sudut pandang:")
BULLET([
    "Berdasarkan output: peringkasan ekstraktif (extractive) yang memilih kalimat dari dokumen asli, dan peringkasan abstraktif (abstractive) yang menghasilkan kalimat baru.",
    "Berdasarkan jumlah dokumen sumber: single-document summarization dan multi-document summarization.",
    "Berdasarkan tujuan: generic summarization (ringkasan umum) dan query-based summarization (ringkasan berdasarkan kata kunci atau pertanyaan).",
    "Berdasarkan domain: domain general (berita, blog, sosial media) dan domain spesifik (medis, hukum, akademik).",
])

H3("2.2.2 Peringkasan Ekstraktif")
P("Peringkasan ekstraktif memilih kalimat-kalimat penting dari dokumen sumber kemudian menggabungkannya menjadi ringkasan baru. Beberapa pendekatan klasik untuk peringkasan ekstraktif adalah:")
BULLET([
    "Frequency-based (Luhn, 1958): kalimat yang banyak mengandung kata-kata berfrekuensi tinggi dianggap penting.",
    "Position-based: kalimat di awal atau akhir paragraf cenderung lebih penting (terutama pada artikel berita).",
    "Length-based: kalimat dengan panjang menengah biasanya lebih informatif.",
    "TF-IDF based (Salton & Buckley, 1988): kalimat dinilai berdasarkan bobot TF-IDF dari kata-kata di dalamnya.",
    "LSA (Latent Semantic Analysis): menggunakan dekomposisi SVD pada matriks term-sentence.",
    "Graph-based / TextRank (Mihalcea & Tarau, 2004): kalimat dijadikan node dalam graph, kemudian diperingkat menggunakan PageRank.",
    "Cluster-based: mengelompokkan kalimat yang mirip lalu memilih representatif dari setiap klaster.",
])
P("Pendekatan TextRank yang berbasis graph PageRank menjadi salah satu metode yang sangat populer karena tidak membutuhkan label/training data dan bekerja secara unsupervised. Selain itu, hasil ringkasan ekstraktif memiliki keunggulan berupa kalimat yang gramatikal karena diambil utuh dari dokumen asli.")

H3("2.2.3 Peringkasan Abstraktif")
P("Peringkasan abstraktif berusaha menghasilkan kalimat baru yang merangkum makna dokumen, tidak harus menggunakan kata atau kalimat yang sama persis dengan dokumen asli. Pendekatan ini meniru cara manusia meringkas, yaitu memahami terlebih dahulu kemudian menulis ulang dengan kata-kata sendiri (paraphrasing).")
P("Sebelum era deep learning, peringkasan abstraktif sangat sulit dilakukan karena memerlukan pemahaman semantik yang mendalam. Setelah munculnya arsitektur Encoder-Decoder dengan mekanisme Attention dan kemudian Transformer (Vaswani et al., 2017), peringkasan abstraktif menjadi lebih mudah dan menghasilkan kualitas yang sangat baik. Beberapa model penting dalam peringkasan abstraktif modern adalah:")
BULLET([
    "Seq2Seq dengan LSTM/GRU + Attention (Bahdanau, 2015).",
    "Transformer (Vaswani et al., 2017) yang menggunakan self-attention sepenuhnya.",
    "Pointer-Generator Network (See et al., 2017) yang dapat menyalin kata dari sumber atau menggenerate kata baru.",
    "BART (Lewis et al., 2020): denoising autoencoder berbasis Transformer.",
    "T5 (Raffel et al., 2020): text-to-text transfer transformer yang menyatukan banyak tugas NLP dalam framework yang sama.",
    "Pegasus (Zhang et al., 2020): pre-training khusus untuk summarization dengan gap-sentence prediction.",
    "mT5 (Xue et al., 2021): versi multilingual T5 yang mendukung 101 bahasa termasuk Indonesia.",
    "XL-Sum (Hasan et al., 2021): hasil fine-tuning mT5 untuk summarization 44 bahasa termasuk Indonesia.",
])

H2("2.3 Term Frequency – Inverse Document Frequency (TF-IDF)")
P("TF-IDF adalah skema pembobotan kata yang sangat banyak digunakan dalam information retrieval dan text mining. Bobot TF-IDF mencerminkan seberapa penting suatu kata pada sebuah dokumen di dalam koleksi dokumen.")
P("Term Frequency (TF) adalah jumlah kemunculan term t dalam dokumen d. Document Frequency (DF) adalah jumlah dokumen yang mengandung term t. Inverse Document Frequency (IDF) adalah logaritma dari rasio total dokumen N terhadap DF, sebagai berikut:")
CODE("TF-IDF(t,d) = TF(t,d) * log( N / (1 + DF(t)) )")
P("Semakin sering sebuah kata muncul di dokumen tertentu (TF tinggi) namun jarang muncul di dokumen lain (DF rendah), maka bobot TF-IDF kata tersebut akan semakin tinggi. Hal ini menandakan bahwa kata tersebut bersifat diskriminatif dan menggambarkan topik dokumen.")

H2("2.4 Cosine Similarity")
P("Cosine similarity mengukur kemiripan antara dua vektor berdasarkan sudut yang dibentuk antara keduanya, bukan berdasarkan magnitude/panjangnya. Nilai cosine similarity berada pada rentang [-1, 1] untuk vektor umum, atau [0, 1] untuk vektor non-negatif (seperti TF-IDF).")
CODE("sim(u, v) = dot(u, v) / ( ||u|| * ||v|| )")
P("Pada peringkasan teks ekstraktif berbasis graph, cosine similarity dipakai untuk mengukur kemiripan antar kalimat (vektor TF-IDF). Hasil ini kemudian dipakai sebagai bobot edge dalam graph yang akan diperingkat oleh PageRank.")

H2("2.5 PageRank dan TextRank")
P("PageRank adalah algoritma yang dikembangkan oleh Larry Page dan Sergey Brin (pendiri Google) untuk meranking halaman web pada hasil pencarian. Ide dasarnya adalah halaman yang banyak ditautkan oleh halaman lain yang penting akan dianggap penting pula. PageRank dihitung secara iteratif dengan rumus:")
CODE("S(Vi) = (1 - d) + d * sum( w_ji / sum_k(w_jk) * S(Vj) )")
P("Di mana d adalah damping factor (umumnya 0,85), Vi adalah node, dan w_ji adalah bobot edge dari Vj ke Vi. TextRank (Mihalcea & Tarau, 2004) mengadopsi konsep PageRank ke domain teks. Setiap kalimat menjadi node dalam graph, dan bobot edge dihitung berdasarkan kemiripan antar kalimat (cosine similarity TF-IDF). Setelah PageRank dijalankan, kalimat-kalimat dengan skor tertinggi dipilih sebagai ringkasan.")

H2("2.6 Arsitektur Transformer")
P("Transformer adalah arsitektur jaringan saraf yang diperkenalkan oleh Vaswani et al. (2017) dalam paper berjudul 'Attention Is All You Need'. Arsitektur ini menggantikan layer recurrent (RNN/LSTM/GRU) dengan mekanisme self-attention yang memungkinkan model memproses seluruh urutan token secara paralel.")
P("Komponen utama Transformer meliputi: (1) Multi-Head Self-Attention yang memungkinkan model memperhatikan berbagai bagian input secara bersamaan; (2) Positional Encoding untuk menyimpan informasi urutan token; (3) Feed-Forward Network di setiap layer; (4) Layer Normalization dan Residual Connection untuk stabilitas training.")
P("Transformer terdiri dari Encoder dan Decoder. Encoder bertugas mengubah input sequence menjadi representasi semantik, sedangkan Decoder menghasilkan output sequence berdasarkan representasi dari Encoder. Untuk tugas summarization, encoder memproses dokumen sumber sedangkan decoder menghasilkan ringkasan secara autoregressive (token demi token).")

H2("2.7 Model mT5 dan XL-Sum")
P("mT5 (Multilingual T5) adalah versi multilingual dari model T5 yang dikembangkan oleh Google Research. mT5 dilatih pada dataset mC4 (multilingual Common Crawl) yang mencakup 101 bahasa, termasuk bahasa Indonesia. Tersedia berbagai ukuran mT5: small (300 juta parameter), base (582 juta), large (1,2 miliar), XL (3,7 miliar), dan XXL (13 miliar).")
P("XL-Sum (Hasan et al., 2021) adalah dataset multilingual untuk summarization berisi 1,35 juta pasangan artikel-ringkasan dalam 44 bahasa, termasuk Indonesia. Model csebuetnlp/mT5_multilingual_XLSum yang digunakan dalam penelitian ini adalah mT5-base yang telah di-fine-tune pada dataset XL-Sum, sehingga dapat langsung dipakai untuk tugas summarization tanpa training tambahan.")

H2("2.8 Metrik Evaluasi ROUGE")
P("ROUGE (Recall-Oriented Understudy for Gisting Evaluation) adalah serangkaian metrik yang diperkenalkan oleh Lin (2004) untuk mengevaluasi kualitas ringkasan otomatis dengan membandingkan output sistem dengan satu atau lebih ringkasan referensi (gold standard). ROUGE saat ini menjadi metrik standar de-facto untuk tugas summarization.")
H3("2.8.1 ROUGE-N")
P("ROUGE-N mengukur overlap n-gram antara ringkasan kandidat dan referensi. ROUGE-1 menghitung overlap unigram (kata tunggal), sedangkan ROUGE-2 menghitung overlap bigram (pasangan dua kata berurutan).")
H3("2.8.2 ROUGE-L")
P("ROUGE-L menggunakan Longest Common Subsequence (LCS) antara dua ringkasan untuk mengukur kemiripan struktural. Berbeda dengan ROUGE-N, ROUGE-L tidak memerlukan kata berurutan secara persis, hanya berurutan secara umum.")
H3("2.8.3 Precision, Recall, dan F1")
P("Setiap varian ROUGE melaporkan tiga nilai: Precision (proporsi n-gram di kandidat yang juga muncul di referensi), Recall (proporsi n-gram di referensi yang juga muncul di kandidat), dan F1 (harmonic mean dari Precision dan Recall):")
CODE("F1 = 2 * (P * R) / (P + R)")

H2("2.9 Penelitian Terkait")
P("Beberapa penelitian terdahulu yang relevan dengan topik ini dapat dilihat pada Tabel 2.1.")
CAPTION("Tabel 2.1 Ringkasan Penelitian Terkait")
TABLE(
    ["No", "Penulis", "Tahun", "Fokus", "Bahasa"],
    [
        ["1", "Luhn", "1958", "Frequency-based summarization", "Inggris"],
        ["2", "Salton & Buckley", "1988", "Term-weighting / TF-IDF", "Inggris"],
        ["3", "Mihalcea & Tarau", "2004", "TextRank graph-based", "Inggris"],
        ["4", "Lin", "2004", "Metrik ROUGE", "Universal"],
        ["5", "Vaswani et al.", "2017", "Transformer architecture", "Universal"],
        ["6", "Lewis et al.", "2020", "BART model", "Inggris"],
        ["7", "Raffel et al.", "2020", "T5 model", "Inggris"],
        ["8", "Zhang et al.", "2020", "Pegasus model", "Inggris"],
        ["9", "Hasan et al.", "2021", "XL-Sum (mT5) 44 bahasa", "Indonesia (incl.)"],
        ["10", "Koto et al.", "2020", "IndoBERT, IndoLEM", "Indonesia"],
        ["11", "Wilie et al.", "2020", "IndoNLU benchmark", "Indonesia"],
    ],
)

H2("2.10 Posisi Penelitian")
P("Penelitian ini menempatkan diri sebagai studi komparatif head-to-head antara metode ekstraktif klasik (TF-IDF + TextRank) dan metode abstraktif modern berbasis pre-trained model (mT5 XL-Sum) pada domain dokumen akademik berbahasa Indonesia. Selain itu, penelitian ini juga menambahkan baseline ketiga berupa LLM untuk subset data sebagai pembanding kualitatif. Implementasi praktis berupa aplikasi web menjadi nilai tambah yang membedakan penelitian ini dari penelitian sebelumnya yang umumnya berhenti di laporan akademik.")

doc.add_page_break()

# ============================================================
# BAB 3 — METODOLOGI
# ============================================================
H1("BAB 3\nMetodologi Penelitian")

H2("3.1 Pendekatan Penelitian")
P("Penelitian ini menggunakan pendekatan kuantitatif eksperimen. Pendekatan kuantitatif dipilih karena hasil yang ingin diukur bersifat numerik (skor ROUGE) dan dapat dihitung secara objektif. Eksperimen dilakukan dengan menjalankan dua metode peringkasan teks yang berbeda pada dataset yang sama, kemudian membandingkan hasilnya menggunakan metrik yang sama.")

H2("3.2 Alur Penelitian")
P("Alur penelitian secara keseluruhan terdiri dari sembilan tahapan utama yang dilakukan secara sekuensial:")
BULLET([
    "Studi pustaka untuk memahami konsep peringkasan teks, NLP bahasa Indonesia, dan model deep learning.",
    "Pengumpulan dataset 100 dokumen akademik berbahasa Indonesia.",
    "Implementasi pipeline preprocessing.",
    "Implementasi metode ekstraktif (TF-IDF + TextRank).",
    "Implementasi metode abstraktif (mT5 multilingual XL-Sum).",
    "Evaluasi kedua metode menggunakan metrik ROUGE.",
    "Perbandingan dengan output Large Language Model (LLM) pada subset 10 dokumen.",
    "Pengembangan aplikasi web interaktif berbasis Flask.",
    "Analisis hasil dan penulisan laporan.",
], numbered=True)

H2("3.3 Dataset")
H3("3.3.1 Sumber Dataset")
P("Dataset penelitian dikumpulkan dari berbagai jurnal akademik dan portal artikel ilmiah berbahasa Indonesia. Pengumpulan dataset dilakukan dengan teknik web scraping menggunakan script scrape_journals.py. Setiap entri dataset terdiri dari tiga kolom utama: judul artikel, isi artikel, dan ringkasan referensi (umumnya berasal dari abstrak artikel). Dataset disimpan pada file data/raw/dataset.csv.")
H3("3.3.2 Statistik Dataset")
CAPTION("Tabel 3.1 Statistik Dataset")
TABLE(
    ["Statistik", "Nilai"],
    [
        ["Jumlah dokumen", "100"],
        ["Bahasa", "Indonesia"],
        ["Domain", "Jurnal akademik"],
        ["Rata-rata panjang dokumen", "25.267 karakter"],
        ["Dokumen terpendek", "1.136 karakter"],
        ["Dokumen terpanjang", "57.745 karakter"],
        ["Rata-rata panjang ringkasan referensi", "964 karakter"],
        ["Format file", "CSV (UTF-8)"],
        ["Sumber", "Scraping jurnal Indonesia"],
    ],
)

H2("3.4 Pipeline Preprocessing")
P("Pipeline preprocessing diimplementasikan pada file src/preprocessor.py. Pipeline ini terdiri dari enam tahapan yang dilakukan secara berurutan:")
BULLET([
    "Case folding: seluruh teks dikonversi menjadi huruf kecil menggunakan method text.lower() pada Python.",
    "Cleaning: penghapusan karakter non-alfabet, URL, mention, hashtag, dan tanda baca berlebih menggunakan regex.",
    "Sentence tokenization: pemecahan paragraf menjadi kalimat menggunakan NLTK sent_tokenize.",
    "Word tokenization: pemecahan kalimat menjadi kata menggunakan NLTK word_tokenize.",
    "Stopword removal: penghapusan stopword bahasa Indonesia menggunakan daftar dari library Sastrawi.",
    "Stemming: pengubahan kata ke bentuk dasar menggunakan Sastrawi Stemmer (Nazief–Adriani).",
], numbered=True)
P("Output preprocessing disimpan ke file output/results/preprocess.json untuk dipakai pada tahap berikutnya.")

H2("3.5 Metode Ekstraktif: TF-IDF + TextRank")
H3("3.5.1 Algoritma")
P("Metode ekstraktif diimplementasikan pada file src/extractive_model.py. Algoritma yang digunakan adalah kombinasi TF-IDF untuk representasi vektor kalimat dan TextRank untuk peringkatan kalimat berdasarkan graph PageRank. Tahapan algoritma adalah sebagai berikut:")
BULLET([
    "Pecah dokumen menjadi himpunan kalimat S = {s1, s2, ..., sn}.",
    "Bangun matriks TF-IDF M berukuran n x v, di mana v adalah ukuran vocabulary.",
    "Hitung matriks similarity W berukuran n x n di mana W_ij = cosine_similarity(Mi, Mj).",
    "Bangun graph berbobot G = (V, E, W) di mana V adalah kalimat dan W adalah matriks similarity.",
    "Jalankan PageRank pada graph G menggunakan NetworkX (nx.pagerank).",
    "Pilih top-k kalimat dengan skor PageRank tertinggi (default k = 5 atau ratio kompresi 20%).",
    "Susun ulang kalimat-kalimat tersebut sesuai urutan kemunculan asli di dokumen.",
], numbered=True)
H3("3.5.2 Hyperparameter")
CAPTION("Tabel 3.2 Hyperparameter Metode Ekstraktif")
TABLE(
    ["Parameter", "Nilai"],
    [
        ["Damping factor (d)", "0,85"],
        ["Maks iterasi PageRank", "100"],
        ["Toleransi konvergensi", "1e-6"],
        ["Vectorizer", "TF-IDF (sklearn)"],
        ["Similarity metric", "Cosine similarity"],
        ["Default top-k", "5"],
        ["Compression ratio", "20%"],
    ],
)

H2("3.6 Metode Abstraktif: mT5 Multilingual XL-Sum")
H3("3.6.1 Model")
P("Metode abstraktif menggunakan model pre-trained csebuetnlp/mT5_multilingual_XLSum yang diunduh dari HuggingFace Model Hub. Model ini adalah varian mT5-base (582 juta parameter) yang telah di-fine-tune pada dataset XL-Sum mencakup 44 bahasa termasuk Indonesia. Implementasi dilakukan pada file src/abstractive_model.py menggunakan library transformers.")
H3("3.6.2 Konfigurasi")
CAPTION("Tabel 3.3 Konfigurasi Model Abstraktif")
TABLE(
    ["Parameter", "Nilai"],
    [
        ["Model", "csebuetnlp/mT5_multilingual_XLSum"],
        ["Arsitektur", "mT5-base, encoder-decoder Transformer"],
        ["Jumlah parameter", "582 juta"],
        ["Tokenizer", "SentencePiece"],
        ["Max source length", "256 token"],
        ["Max target length", "128 token"],
        ["Min target length", "30 token"],
        ["Beam search", "4 beams"],
        ["Length penalty", "1,0"],
        ["No-repeat n-gram size", "2"],
        ["Inferensi", "CPU (AMD EPYC 7763)"],
    ],
)

H2("3.7 Evaluasi: Metrik ROUGE")
P("Evaluasi dilakukan menggunakan library rouge-score dari Google. Untuk setiap pasangan ringkasan kandidat dan ringkasan referensi, dihitung tiga varian ROUGE: ROUGE-1 (unigram overlap), ROUGE-2 (bigram overlap), dan ROUGE-L (longest common subsequence). Setiap varian melaporkan Precision, Recall, dan F1.")
P("Skor agregat dihitung dengan rata-rata aritmatik atas seluruh dokumen pada dataset. Hasil evaluasi disimpan pada file output/results/evaluate.json.")

H2("3.8 Perbandingan dengan LLM")
P("Sebagai third baseline, dilakukan perbandingan dengan ringkasan yang dihasilkan oleh Large Language Model (LLM). Perbandingan ini dilakukan pada subset 10 dokumen karena keterbatasan biaya API dan waktu. Tiga ringkasan dihasilkan untuk setiap dokumen: ringkasan ekstraktif, ringkasan abstraktif, dan ringkasan LLM. Skor ROUGE dihitung untuk semuanya. Hasil disimpan pada file output/results/comparison_summary.json.")

H2("3.9 Aplikasi Web")
H3("3.9.1 Stack Teknologi")
BULLET([
    "Backend: Flask (Python web framework).",
    "Templating: Jinja2.",
    "Frontend: Bootstrap 5 + Chart.js untuk visualisasi.",
    "Server produksi: Gunicorn + Nginx.",
    "Storage: file CSV dan JSON.",
])
H3("3.9.2 Halaman Utama")
BULLET([
    "/ : input dokumen (CSV upload, PDF upload, atau teks manual).",
    "/summarize : menjalankan kedua metode dan menampilkan hasil ringkasan.",
    "/evaluate : menjalankan ROUGE dan menampilkan skor.",
    "/analysis : dashboard chart (line, bar, scatter, distribusi histogram).",
    "/detail/<id> : detail per-dokumen (ringkasan, referensi, skor ROUGE).",
])
H3("3.9.3 Deployment")
P("Aplikasi di-deploy menggunakan Gunicorn sebagai application server dan Nginx sebagai reverse proxy. Konfigurasi terdapat pada folder deploy/ (file nlp-summarization.service untuk systemd, dan nginx-nlp-summarization.conf untuk Nginx).")

H2("3.10 Spesifikasi Sistem")
CAPTION("Tabel 3.4 Spesifikasi Sistem")
TABLE(
    ["Komponen", "Spesifikasi"],
    [
        ["CPU", "AMD EPYC 7763 (Codespaces dev container)"],
        ["RAM", "Cukup untuk inferensi mT5-base"],
        ["GPU", "Tidak digunakan (CPU-only)"],
        ["OS", "Ubuntu 24.04.3 LTS"],
        ["Bahasa pemrograman", "Python 3.x"],
        ["Library NLP", "Sastrawi, NLTK, scikit-learn, NetworkX"],
        ["Library deep learning", "Transformers (HuggingFace), PyTorch"],
        ["Library evaluasi", "rouge-score"],
        ["Library web", "Flask, Jinja2, Gunicorn"],
        ["Library dokumen", "python-docx"],
    ],
)

doc.add_page_break()

# ============================================================
# BAB 4 — HASIL DAN PEMBAHASAN
# ============================================================
H1("BAB 4\nHasil dan Pembahasan")

H2("4.1 Hasil Preprocessing")
P("Pipeline preprocessing berhasil dijalankan pada seluruh 100 dokumen tanpa error. Output preprocessing meliputi teks hasil case folding, hasil cleaning, hasil tokenization (kalimat dan kata), hasil stopword removal, dan hasil stemming. Output disimpan pada file output/results/preprocess.json.")
P("Sebagai gambaran, satu dokumen awal yang berisi 25.267 karakter dapat dipecah menjadi sekitar 200–400 kalimat dan ribuan token. Setelah stopword removal dan stemming, jumlah token unik berkurang signifikan, sehingga matriks TF-IDF yang dihasilkan menjadi lebih ringkas.")

H2("4.2 Hasil Ringkasan")
P("Ringkasan ekstraktif untuk seluruh 100 dokumen tersimpan pada file output/summaries/extractive_summaries.csv, sedangkan ringkasan abstraktif tersimpan pada file output/summaries/abstractive_summaries.csv. Setiap baris file CSV memuat doc_id, judul, ringkasan referensi, dan ringkasan hasil sistem.")
P("Secara kualitatif, ringkasan ekstraktif cenderung lebih panjang karena memilih kalimat-kalimat utuh dari dokumen asli, sementara ringkasan abstraktif lebih pendek dan padat karena dibatasi oleh max target length 128 token.")

H2("4.3 Hasil Evaluasi ROUGE — Dataset Penuh (100 dokumen)")
ext = evaluate["extractive_scores"]
abs_ = evaluate["abstractive_scores"]
CAPTION("Tabel 4.1 Skor ROUGE Agregat (100 dokumen)")
TABLE(
    ["Metrik", "Metode", "Precision", "Recall", "F1-Score"],
    [
        ["ROUGE-1", "Ekstraktif", f"{ext['rouge1']['precision']:.4f}", f"{ext['rouge1']['recall']:.4f}", f"{ext['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-1", "Abstraktif", f"{abs_['rouge1']['precision']:.4f}", f"{abs_['rouge1']['recall']:.4f}", f"{abs_['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-2", "Ekstraktif", f"{ext['rouge2']['precision']:.4f}", f"{ext['rouge2']['recall']:.4f}", f"{ext['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-2", "Abstraktif", f"{abs_['rouge2']['precision']:.4f}", f"{abs_['rouge2']['recall']:.4f}", f"{abs_['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-L", "Ekstraktif", f"{ext['rougeL']['precision']:.4f}", f"{ext['rougeL']['recall']:.4f}", f"{ext['rougeL']['fmeasure']:.4f}"],
        ["ROUGE-L", "Abstraktif", f"{abs_['rougeL']['precision']:.4f}", f"{abs_['rougeL']['recall']:.4f}", f"{abs_['rougeL']['fmeasure']:.4f}"],
    ],
)
P(f"Pada Tabel 4.1 dapat dilihat bahwa metode ekstraktif unggul pada F1-score di seluruh varian ROUGE: ROUGE-1 sebesar {ext['rouge1']['fmeasure']:.4f} (vs {abs_['rouge1']['fmeasure']:.4f} abstraktif), ROUGE-2 sebesar {ext['rouge2']['fmeasure']:.4f} (vs {abs_['rouge2']['fmeasure']:.4f}), dan ROUGE-L sebesar {ext['rougeL']['fmeasure']:.4f} (vs {abs_['rougeL']['fmeasure']:.4f}).")
P(f"Recall metode ekstraktif jauh lebih tinggi dibandingkan abstraktif. Pada ROUGE-1, recall ekstraktif {ext['rouge1']['recall']:.4f} sedangkan abstraktif hanya {abs_['rouge1']['recall']:.4f}. Hal ini wajar karena metode ekstraktif memilih kalimat-kalimat utuh sehingga banyak token referensi tertangkap. Sebaliknya, precision metode abstraktif jauh lebih tinggi: pada ROUGE-1 precision abstraktif {abs_['rouge1']['precision']:.4f} sedangkan ekstraktif {ext['rouge1']['precision']:.4f}. Hal ini menunjukkan bahwa kata-kata yang dihasilkan mT5 memang relevan dengan referensi, namun jumlahnya terbatas akibat batas 128 token output.")

H2("4.4 Hasil Evaluasi ROUGE — Subset 10 Dokumen + LLM")
ce = comparison["extractive"]
ca = comparison["abstractive"]
CAPTION("Tabel 4.2 Skor ROUGE Subset (10 dokumen, dengan referensi LLM)")
TABLE(
    ["Metrik", "Metode", "Precision", "Recall", "F1-Score"],
    [
        ["ROUGE-1", "Ekstraktif", f"{ce['rouge1']['precision']:.4f}", f"{ce['rouge1']['recall']:.4f}", f"{ce['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-1", "Abstraktif", f"{ca['rouge1']['precision']:.4f}", f"{ca['rouge1']['recall']:.4f}", f"{ca['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-2", "Ekstraktif", f"{ce['rouge2']['precision']:.4f}", f"{ce['rouge2']['recall']:.4f}", f"{ce['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-2", "Abstraktif", f"{ca['rouge2']['precision']:.4f}", f"{ca['rouge2']['recall']:.4f}", f"{ca['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-L", "Ekstraktif", f"{ce['rougeL']['precision']:.4f}", f"{ce['rougeL']['recall']:.4f}", f"{ce['rougeL']['fmeasure']:.4f}"],
        ["ROUGE-L", "Abstraktif", f"{ca['rougeL']['precision']:.4f}", f"{ca['rougeL']['recall']:.4f}", f"{ca['rougeL']['fmeasure']:.4f}"],
    ],
)
P(f"Pada subset 10 dokumen dengan referensi yang juga dihasilkan LLM, hasil menjadi berbalik: metode abstraktif unggul pada seluruh varian ROUGE. ROUGE-1 abstraktif {ca['rouge1']['fmeasure']:.4f} vs ekstraktif {ce['rouge1']['fmeasure']:.4f}; ROUGE-2 abstraktif {ca['rouge2']['fmeasure']:.4f} vs ekstraktif {ce['rouge2']['fmeasure']:.4f}; dan ROUGE-L abstraktif {ca['rougeL']['fmeasure']:.4f} vs ekstraktif {ce['rougeL']['fmeasure']:.4f}.")
P(f"Best Method pada subset ini adalah {comparison['best_method']}. Hal ini terjadi karena pada subset, ringkasan referensi dihasilkan oleh LLM yang juga merupakan model neural berbasis Transformer, sehingga gaya bahasanya lebih dekat dengan output mT5 dibandingkan dengan ringkasan tradisional yang panjang dan ekstraktif.")

H2("4.5 Pembahasan Trade-off Precision vs Recall")
P("Salah satu temuan kunci penelitian ini adalah profil precision-recall yang berkebalikan antara dua metode. Metode ekstraktif memberikan recall tinggi karena memilih kalimat-kalimat utuh dengan banyak kata pendukung. Sebaliknya, metode abstraktif memberikan precision tinggi karena output yang dihasilkan padat dan relevan, tetapi recall menjadi rendah karena cakupan kontennya terbatas oleh max output 128 token.")
P("Dalam konteks praktis, pemilihan metode bergantung pada use-case pengguna. Apabila pengguna menginginkan ringkasan dengan cakupan yang luas, misalnya untuk literature review atau sebagai dasar analisis lebih lanjut, maka metode ekstraktif lebih sesuai. Sebaliknya, apabila pengguna membutuhkan ringkasan yang singkat, padat, dan rapi seperti abstrak singkat, maka metode abstraktif menjadi pilihan yang lebih tepat.")

H2("4.6 Analisis Faktor Per-Dokumen")
P("Analisis per-dokumen dilakukan dengan memeriksa skor ROUGE individual untuk setiap dokumen pada dataset 100 dokumen. Tiga faktor utama yang berkorelasi dengan variasi performa antar dokumen adalah:")
H3("4.6.1 Panjang Dokumen")
P("Dokumen yang sangat panjang (lebih dari 40.000 karakter) cenderung menghasilkan F1 abstraktif yang lebih rendah. Hal ini disebabkan oleh batas max source length 256 token, sehingga sebagian besar isi dokumen terpotong sebelum diproses oleh model. Selain itu, batas max target length 128 token juga membatasi kemampuan model untuk memberikan ringkasan yang komprehensif.")
H3("4.6.2 Jumlah Kalimat")
P("Dokumen dengan jumlah kalimat banyak memberikan keuntungan bagi metode ekstraktif karena tersedia banyak kandidat untuk dipilih oleh TextRank. Sebaliknya, dokumen pendek dengan sedikit kalimat (kurang dari 10) menghasilkan ringkasan ekstraktif yang nyaris identik dengan dokumen aslinya, sehingga skor recall sangat tinggi tetapi precision rendah.")
H3("4.6.3 Rasio Kompresi")
P("Rasio kompresi adalah panjang ringkasan referensi dibagi panjang dokumen. Rasio rendah (ringkasan sangat pendek dari dokumen yang sangat panjang) lebih sulit untuk kedua metode. Namun, metode ekstraktif lebih tahan terhadap rasio rendah karena dapat memilih sub-himpunan kalimat dengan ukuran yang fleksibel.")

H2("4.7 Pembahasan Mengapa Hasil Berbeda per Dokumen")
P("Selain faktor numerik di atas, terdapat juga beberapa faktor kualitatif yang memengaruhi variasi hasil per dokumen:")
BULLET([
    "Struktur dokumen: dokumen dengan kalimat topik (topic sentence) yang jelas pada paragraf pertama akan menguntungkan metode ekstraktif berbasis TextRank karena kalimat topik biasanya memiliki similarity tinggi dengan kalimat lain.",
    "Overlap vocabulary dengan referensi: jika ringkasan referensi banyak menggunakan kata-kata yang sama dengan dokumen sumber, metode ekstraktif akan unggul karena dapat menyalin langsung. Sebaliknya, jika referensi banyak melakukan paraphrasing, metode abstraktif akan lebih unggul.",
    "Gaya referensi: referensi yang ditulis ulang oleh manusia dengan gaya naratif lebih cocok dievaluasi terhadap output abstraktif. Referensi yang berupa kutipan langsung lebih cocok untuk evaluasi ekstraktif.",
    "Domain spesifik: terminologi teknis yang langka kadang tidak ter-generate dengan baik oleh mT5 karena efek Out-Of-Vocabulary (OOV) pada SentencePiece tokenizer. Akibatnya, kata-kata teknis cenderung digantikan dengan kata yang lebih umum.",
])

H2("4.8 Aplikasi Web")
P("Aplikasi web berbasis Flask berhasil dikembangkan dan dapat dijalankan baik secara lokal maupun di server produksi (Gunicorn + Nginx). Fitur-fitur utama aplikasi meliputi:")
BULLET([
    "Multi-format input: pengguna dapat mengunggah file CSV (untuk batch processing), file PDF (akan diekstrak otomatis), atau menempelkan teks manual.",
    "Visualisasi pipeline langkah demi langkah: pengguna dapat melihat hasil pada setiap tahap preprocessing.",
    "Halaman /analysis dengan dashboard berisi line chart skor ROUGE per dokumen, bar chart perbandingan metode, scatter plot panjang dokumen vs F1, dan histogram distribusi.",
    "Perbandingan side-by-side antara ringkasan ekstraktif, abstraktif, dan LLM (jika tersedia).",
    "Antarmuka bilingual (Indonesia/Inggris) untuk memudahkan pengguna internasional.",
    "Halaman detail per-dokumen yang menampilkan dokumen asli, ringkasan referensi, ringkasan kedua metode, dan skor ROUGE.",
])

H2("4.9 Diskusi Keterbatasan")
P("Penelitian ini memiliki beberapa keterbatasan yang perlu dicatat untuk penelitian selanjutnya:")
BULLET([
    "Inferensi mT5 dilakukan pada CPU sehingga waktu eksekusi cukup lama untuk batch besar (kurang lebih beberapa detik per dokumen).",
    "Model mT5 yang digunakan tidak melalui fine-tuning pada domain akademik Indonesia (zero-shot dari XL-Sum yang dilatih pada artikel berita BBC).",
    "Ukuran dataset 100 dokumen cukup untuk evaluasi awal tetapi belum ideal untuk training model atau analisis statistik yang lebih mendalam.",
    "Metrik ROUGE bersifat n-gram dan tidak menangkap kesetaraan semantik. Sebuah paraphrase yang benar secara makna dapat memperoleh skor ROUGE rendah karena tidak ada overlap kata.",
    "Belum dilakukan evaluasi manusia (human evaluation) untuk menilai aspek fluency, coherence, dan faithfulness dari ringkasan.",
    "Dataset hanya berasal dari satu domain (jurnal akademik), sehingga generalisasi hasil ke domain lain (berita, dokumen hukum, dokumen medis) belum dapat dipastikan.",
])

doc.add_page_break()

# ============================================================
# BAB 5 — KESIMPULAN DAN SARAN
# ============================================================
H1("BAB 5\nKesimpulan dan Saran")

H2("5.1 Kesimpulan")
P("Berdasarkan hasil penelitian dan pembahasan pada Bab 4, dapat ditarik kesimpulan sebagai berikut:")
BULLET([
    "Penelitian berhasil mengimplementasikan dan membandingkan dua metode peringkasan teks akademik berbahasa Indonesia, yaitu metode ekstraktif berbasis TF-IDF + TextRank dan metode abstraktif berbasis mT5 multilingual XL-Sum.",
    f"Pada dataset penuh 100 dokumen, metode ekstraktif unggul secara F1-score pada seluruh varian ROUGE: ROUGE-1 F1 sebesar {ext['rouge1']['fmeasure']:.4f} (ekstraktif) berbanding {abs_['rouge1']['fmeasure']:.4f} (abstraktif), ROUGE-2 F1 sebesar {ext['rouge2']['fmeasure']:.4f} berbanding {abs_['rouge2']['fmeasure']:.4f}, dan ROUGE-L F1 sebesar {ext['rougeL']['fmeasure']:.4f} berbanding {abs_['rougeL']['fmeasure']:.4f}.",
    f"Pada subset 10 dokumen dengan referensi yang dihasilkan LLM, metode abstraktif justru unggul: ROUGE-1 F1 sebesar {ca['rouge1']['fmeasure']:.4f} (abstraktif) berbanding {ce['rouge1']['fmeasure']:.4f} (ekstraktif), dan ROUGE-L F1 sebesar {ca['rougeL']['fmeasure']:.4f} berbanding {ce['rougeL']['fmeasure']:.4f}.",
    "Terdapat trade-off konsisten antara precision dan recall: metode ekstraktif memberikan recall tinggi (cakupan luas) sedangkan metode abstraktif memberikan precision tinggi (output padat dan akurat tetapi pendek).",
    "Faktor panjang dokumen, jumlah kalimat, dan rasio kompresi secara signifikan memengaruhi performa per dokumen. Dokumen panjang dan rasio kompresi rendah cenderung menyulitkan metode abstraktif karena batas output 128 token.",
    "Aplikasi web berbasis Flask berhasil dibangun sebagai bentuk implementasi praktis. Aplikasi menyediakan pipeline end-to-end yang interaktif untuk pengguna non-teknis dengan fitur multi-format input, visualisasi, dan dashboard analisis.",
    "Pemilihan metode peringkasan teks dalam praktik bergantung pada use-case: ekstraktif untuk cakupan luas dan abstraktif untuk ringkasan padat.",
], numbered=True)

H2("5.2 Saran")
P("Berdasarkan keterbatasan dan temuan penelitian, beberapa saran untuk penelitian selanjutnya adalah sebagai berikut:")
BULLET([
    "Lakukan fine-tuning model mT5 pada korpus akademik Indonesia menggunakan GPU untuk meningkatkan recall metode abstraktif dan adaptasi domain.",
    "Eksplorasi pendekatan hybrid: jalankan metode ekstraktif terlebih dahulu untuk memilih kalimat-kalimat penting, kemudian gunakan model abstraktif untuk paraphrasing. Pendekatan ini diharapkan dapat meningkatkan precision dan recall sekaligus.",
    "Perbesar ukuran dataset menjadi lebih dari 1.000 dokumen lintas disiplin (teknik, sosial, kedokteran, ekonomi) untuk meningkatkan generalisasi hasil.",
    "Tambahkan metrik evaluasi modern seperti BERTScore, METEOR, dan BLEURT yang dapat menangkap kesetaraan semantik selain n-gram overlap.",
    "Lakukan evaluasi manusia (human evaluation) dengan responden ahli untuk menilai aspek fluency (kelancaran), coherence (kepaduan), dan faithfulness (kesetiaan terhadap dokumen sumber).",
    "Eksplorasi model lain seperti IndoBART, IndoT5, dan model open-source LLM berbahasa Indonesia (Llama-Indo, Sahabat-AI, Komodo) sebagai pembanding tambahan.",
    "Atasi batas panjang output abstraktif dengan strategi chunking + hierarchical summarization untuk dokumen yang sangat panjang (>40.000 karakter).",
    "Tambahkan fitur eksport ringkasan ke format Word/PDF langsung dari aplikasi web.",
    "Pertimbangkan integrasi dengan sistem manajemen referensi (Mendeley, Zotero) untuk meningkatkan utilitas aplikasi.",
    "Lakukan studi user experience (UX) untuk mengukur seberapa berguna aplikasi bagi pengguna riil di lingkungan akademik.",
], numbered=True)

H2("5.3 Penutup")
P("Demikian laporan penelitian Analisis Komparatif Metode Peringkasan Teks Ekstraktif dan Abstraktif untuk Dokumen Akademik Berbahasa Indonesia ini disusun. Penulis menyadari bahwa penelitian ini masih memiliki banyak ruang untuk perbaikan, baik dari segi metodologi, dataset, maupun implementasi. Penulis berharap penelitian ini dapat menjadi referensi awal bagi peneliti, akademisi, dan praktisi yang ingin mengembangkan sistem peringkasan teks berbahasa Indonesia di masa mendatang. Kritik dan saran yang membangun sangat penulis harapkan untuk kesempurnaan penelitian selanjutnya.")

out_path = OUT_DIR / "skripsi_bab1-5.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {out_path.stat().st_size / 1024:.1f} KB")
